from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

def docx_to_html(doc_file):
    doc = Document(doc_file)
    html = []

    for para in doc.paragraphs:
        text = para.text
        if text:
            runs = para.runs
            if runs:
                run_html = []
                for run in runs:
                    run_text = run.text
                    if run_text:
                        run_html.append(run_text)
                    font = run.font
                    if font:
                        style = ""
                        if font.bold:
                            style += "font-weight:bold;"
                        if font.italic:
                            style += "font-style:italic;"
                        if font.underline:
                            style += "text-decoration:underline;"
                        color = font.color.rgb
                        if color:
                            r, g, b = color
                            style += f"color:rgb({r}, {g}, {b});"
                        size = font.size.pt
                        if size:
                            style += f"font-size:{size}pt;"
                        if style:
                            run_html[-1] = f"<span style='{style}'>{run_text}</span>"
                text = "".join(run_html)
            align = para.alignment
            if align != WD_PARAGRAPH_ALIGNMENT.LEFT:
                html.append(f"<p align='{align._member_name.lower()}'>{text}</p>")
            else:
                html.append(f"<p>{text}</p>")

    for table in doc.tables:
        rows = table.rows
        cols = table.columns
        html.append("<table>")
        for row in rows:
            html.append("<tr>")
            for cell in row.cells:
                cell_html = []
                cell_text = cell.text
                if cell_text:
                    cell_html.append(cell_text)
                shading_color = cell.shading.background_color.rgb
                if shading_color:
                    r, g, b = shading_color
                    cell_html.insert(0, f"<div style='background-color:rgb({r}, {g}, {b});'>")
                    cell_html.append("</div>")
                cell_text = "".join(cell_html)
                cell_width = cell.width
                if cell_width:
                    html.append(f"<td width='{cell_width}'>{cell_text}</td>")
                else:
                    html.append(f"<td>{cell_text}</td>")
            html.append("</tr>")
        html.append("</table>")

    return "".join(html)
