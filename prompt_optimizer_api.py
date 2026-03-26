"""
SIPDO Prompt Optimizer Pipeline

Closed-loop prompt optimization using Anthropic Claude agents and GPT evaluation.
Implements the SIPDO framework (Self-Improving Prompts through Data-Augmented
Optimization) adapted for document field extraction.

Pipeline Steps:
    1. Document Analysis          — understand document structure
    2. Field Decomposition        — break complex fields into helpers
    3. Seed Prompt Generation     — create initial extraction prompt
    4. SIPDO Optimization Loop    — up to 5 iterations of:
       a. Synthetic data generation (progressive difficulty)
       b. Prompt evaluation via GPT (direct_llm or coreplus)
       c. Error analysis + prompt refinement
       d. Local + global verification
    5. Consistency Test           — final validation
    6. Packaging                  — return optimized prompt + artifacts

Reference: https://arxiv.org/abs/2505.19514
Author: CoreRAG Team
Version: 1.0.0
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

import io
import json
import logging
import time
import uuid
import re
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple

import yaml
import PyPDF2
try:
    import docx
except ImportError:
    docx = None

from logger import get_logger
from agents import load_domain_knowledge
from apis.prompt_optimizer_agents import (
    TokenTracker,
    agent_analyze_document,
    agent_decompose_fields,
    agent_generate_seed_prompt,
    agent_generate_synthetic_data,
    agent_analyze_errors,
    agent_refine_prompt,
    agent_audit_consistency,
    evaluate_prompt_with_gpt,
)

logger = get_logger(__name__)


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

def _load_optimizer_config() -> Dict[str, Any]:
    """Load prompt optimizer config from dev.yaml."""
    config_path = Path(__file__).parent.parent / "dev.yaml"
    try:
        with open(config_path, "r") as f:
            full_cfg = yaml.safe_load(f)
        return full_cfg.get("prompt_optimizer", {})
    except Exception:
        return {}


_OPT_CFG = _load_optimizer_config()

MAX_ITERATIONS = _OPT_CFG.get("max_iterations", 5)
MAX_DIFFICULTY = _OPT_CFG.get("max_difficulty", 5)
ANTHROPIC_MODEL = _OPT_CFG.get("anthropic_model", "claude-sonnet-4-20250514")
DEFAULT_EVAL_MODE = _OPT_CFG.get("default_evaluation_mode", "direct_llm")


def _get_anthropic_key() -> str:
    """Get Anthropic API key from dev.yaml."""
    config_path = Path(__file__).parent.parent / "dev.yaml"
    try:
        with open(config_path, "r") as f:
            full_cfg = yaml.safe_load(f)
        return full_cfg.get("anthropic", {}).get("api_key", "")
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# Per-session SIPDO Audit Logger
# ---------------------------------------------------------------------------

def _create_audit_logger(pipeline_id: str, filename: str) -> logging.Logger:
    """Create a per-session file logger that writes to logs/sipdo_audit_<id>.log."""
    logs_dir = Path(__file__).parent.parent / "logs"
    logs_dir.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    safe_name = re.sub(r"[^\w\-.]", "_", filename)
    log_path = logs_dir / f"sipdo_audit_{safe_name}_{ts}_{pipeline_id}.log"

    audit = logging.getLogger(f"sipdo_audit.{pipeline_id}")
    audit.setLevel(logging.DEBUG)
    audit.propagate = False
    # Avoid duplicate handlers on repeated calls
    if not audit.handlers:
        fh = logging.FileHandler(str(log_path), encoding="utf-8")
        fh.setFormatter(logging.Formatter("%(asctime)s | %(message)s", datefmt="%Y-%m-%d %H:%M:%S"))
        audit.addHandler(fh)
    logger.info(f"📋 SIPDO audit log → {log_path}")
    return audit


# ---------------------------------------------------------------------------
# Document Text Extraction (reuse DTE patterns)
# ---------------------------------------------------------------------------

def extract_text_from_pdf(pdf_content: bytes, filename: str) -> str:
    """Extract full text from PDF bytes."""
    try:
        pdf_file = io.BytesIO(pdf_content)
        reader = PyPDF2.PdfReader(pdf_file)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(f"[Page {i + 1}]\n{text}")
        full_text = "\n\n".join(pages)
        logger.info(f"Extracted {len(reader.pages)} pages from PDF: {filename}")
        return full_text
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise


def extract_text_from_docx(docx_content: bytes, filename: str) -> str:
    """Extract full text from DOCX bytes."""
    if docx is None:
        raise ImportError("python-docx is required for DOCX processing")
    try:
        docx_file = io.BytesIO(docx_content)
        document = docx.Document(docx_file)
        paragraphs = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        full_text = "\n\n".join(paragraphs)
        logger.info(f"Extracted {len(paragraphs)} paragraphs from DOCX: {filename}")
        return full_text
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise


def extract_document_text(file_content: bytes, filename: str) -> str:
    """Route to PDF or DOCX extractor based on extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return extract_text_from_pdf(file_content, filename)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_content, filename)
    else:
        raise ValueError(f"Unsupported file type '.{ext}'. Supported: .pdf, .docx")


# ---------------------------------------------------------------------------
# Field Comparison Utilities
# ---------------------------------------------------------------------------

def _normalize_value(value: Any) -> str:
    """Normalize a value for comparison."""
    if value is None:
        return ""
    s = str(value).strip().lower()
    # Remove common formatting differences
    s = re.sub(r"[\s,]+", " ", s)
    s = re.sub(r"^\$\s*", "", s)
    return s


def _compare_field(expected: str, actual: Any) -> bool:
    """Compare expected vs actual field value with fuzzy matching."""
    if actual is None:
        return False
    norm_exp = _normalize_value(expected)
    norm_act = _normalize_value(actual)
    if not norm_exp:
        return True  # No expected value to compare
    # Exact match after normalization
    if norm_exp == norm_act:
        return True
    # Containment check (actual contains expected or vice versa)
    if norm_exp in norm_act or norm_act in norm_exp:
        return True
    return False


def _compute_accuracy(
    expected_values: Dict[str, str],
    actual_values: Dict[str, Any],
) -> Tuple[float, List[str], Dict[str, Dict[str, str]]]:
    """Compute accuracy and identify failed fields.

    Returns: (accuracy, failed_field_names, per_field_details)
    """
    total = len(expected_values)
    if total == 0:
        return 1.0, [], {}
    correct = 0
    failed = []
    details = {}
    for field_name, expected in expected_values.items():
        actual = actual_values.get(field_name)
        if _compare_field(expected, actual):
            correct += 1
            details[field_name] = {"status": "pass", "expected": expected, "actual": str(actual or "")}
        else:
            failed.append(field_name)
            details[field_name] = {"status": "fail", "expected": expected, "actual": str(actual or "")}
    return correct / total, failed, details


# ---------------------------------------------------------------------------
# CorePlus Evaluation Mode
# ---------------------------------------------------------------------------

def _evaluate_via_coreplus(
    prompt_set: Dict[str, Any],
    file_content: bytes,
    filename: str,
) -> Dict[str, Any]:
    """Evaluate by calling the /pipeline_coreplus endpoint internally."""
    import requests

    api_url = "http://localhost:8000/pipeline_coreplus"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "pdf"
    mime = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if ext == "docx"
        else "application/pdf"
    )
    files = {"file": (filename, file_content, mime)}
    data = {
        "prompt_set": json.dumps(prompt_set),
        "mode": "Auto",
        "postprocessing": "on",
    }
    try:
        resp = requests.post(api_url, files=files, data=data, timeout=1800)
        if resp.status_code == 200:
            result = resp.json()
            return result.get("extracted_fields", {})
        else:
            logger.error(f"CorePlus evaluation failed: HTTP {resp.status_code}")
            return {}
    except Exception as e:
        logger.error(f"CorePlus evaluation error: {e}")
        return {}


# ---------------------------------------------------------------------------
# Main Pipeline Service
# ---------------------------------------------------------------------------

class PromptOptimizerService:
    """
    SIPDO-based Prompt Optimizer Service.

    Orchestrates the closed-loop optimization of extraction prompts using
    Anthropic Claude agents for reasoning and GPT for evaluation.
    """

    def __init__(self):
        logger.info("=" * 80)
        logger.info("INITIALIZING PROMPT OPTIMIZER SERVICE")
        logger.info("=" * 80)
        self.anthropic_key = _get_anthropic_key()
        if not self.anthropic_key:
            logger.warning("⚠️  No Anthropic API key configured in dev.yaml")
        logger.info("✅ PROMPT OPTIMIZER SERVICE READY")
        logger.info("=" * 80)

    def optimize_prompt(
        self,
        file_content: bytes,
        filename: str,
        expected_values: Dict[str, str],
        max_iterations: int = 5,
        difficulty_levels: int = 5,
        evaluation_mode: str = "direct_llm",
        field_decomposition: bool = True,
        consistency_test: bool = True,
    ) -> Dict[str, Any]:
        """
        Main entry point: optimize an extraction prompt using the SIPDO loop.

        Args:
            file_content: Raw PDF/DOCX bytes
            filename: Original filename
            expected_values: {field_name: expected_value} pairs
            max_iterations: Max SIPDO loop iterations (1-5)
            difficulty_levels: Max difficulty for synthetic data (1-5)
            evaluation_mode: "direct_llm" or "coreplus"
            field_decomposition: Whether to auto-decompose complex fields
            consistency_test: Whether to run final consistency audit

        Returns:
            Full result dict with optimized prompt, trace, reports, etc.
        """
        pipeline_id = str(uuid.uuid4())[:8]
        start_time = time.time()
        tracker = TokenTracker()
        optimization_trace = []
        audit = _create_audit_logger(pipeline_id, filename)

        # Clamp parameters
        max_iterations = min(max(max_iterations, 1), MAX_ITERATIONS)
        difficulty_levels = min(max(difficulty_levels, 1), MAX_DIFFICULTY)

        logger.info("=" * 80)
        logger.info(f"PROMPT OPTIMIZER STARTED — ID: {pipeline_id}")
        logger.info(f"  Document       : {filename}")
        logger.info(f"  Fields         : {len(expected_values)}")
        logger.info(f"  Max iterations : {max_iterations}")
        logger.info(f"  Difficulty     : {difficulty_levels}")
        logger.info(f"  Eval mode      : {evaluation_mode}")
        logger.info("=" * 80)

        audit.info("=" * 80)
        audit.info(f"SIPDO PROMPT OPTIMIZER AUDIT — Pipeline {pipeline_id}")
        audit.info(f"Document: {filename}")
        audit.info(f"Fields: {json.dumps(list(expected_values.keys()))}")
        audit.info(f"Expected values: {json.dumps(expected_values, indent=2)}")
        audit.info(f"Max iterations: {max_iterations} | Difficulty: {difficulty_levels} | Eval mode: {evaluation_mode}")
        audit.info("=" * 80)

        api_key = self.anthropic_key
        if not api_key:
            return {
                "status": "error",
                "pipeline_id": pipeline_id,
                "error": "No Anthropic API key configured. Add anthropic.api_key to dev.yaml.",
                "processing_time_seconds": time.time() - start_time,
                "timestamp": datetime.utcnow().isoformat(),
            }

        try:
            # ── STEP 1: Extract document text ─────────────────────────
            logger.info("\n📄 STEP 1: Document Text Extraction")
            t0 = time.time()
            document_text = extract_document_text(file_content, filename)
            if not document_text.strip():
                raise ValueError("No text could be extracted from the document")
            step1_time = time.time() - t0
            logger.info(f"   ✅ Extracted {len(document_text)} chars in {step1_time:.1f}s")

            # ── STEP 1b: Load Domain Knowledge ───────────────────────
            # Use document text + field names to match relevant domain knowledge
            domain_query = document_text[:3000] + " " + " ".join(expected_values.keys())
            domain_knowledge = load_domain_knowledge(domain_query)
            if domain_knowledge:
                logger.info(f"   📚 Loaded domain knowledge ({len(domain_knowledge)} chars)")
            else:
                logger.info("   📚 No matching domain knowledge found")

            # ── STEP 2: Document Analysis (Claude) ────────────────────
            logger.info("\n🔍 STEP 2: Document Analysis")
            t0 = time.time()
            analysis_result = agent_analyze_document(
                document_text, expected_values, api_key, tracker,
                domain_knowledge=domain_knowledge,
            )
            if not analysis_result.success:
                raise RuntimeError(f"Document analysis failed: {analysis_result.error}")
            document_analysis = analysis_result.parsed_json or {}
            step2_time = time.time() - t0
            logger.info(f"   ✅ Document analysis complete in {step2_time:.1f}s")

            # ── STEP 3: Field Decomposition (Claude) ──────────────────
            helper_fields = {}
            decomposition_report = {}
            if field_decomposition:
                logger.info("\n🔧 STEP 3: Field Decomposition")
                t0 = time.time()
                decomp_result = agent_decompose_fields(
                    expected_values, document_analysis, api_key, tracker
                )
                if decomp_result.success and decomp_result.parsed_json:
                    decomposition_report = decomp_result.parsed_json
                    # Extract helper field info
                    for fname, decomp in decomposition_report.get("decompositions", {}).items():
                        helper_fields[fname] = {
                            "sub_fields": list(decomp.get("sub_fields", {}).keys()),
                            "python_postprocessing": decomp.get("python_postprocessing", ""),
                        }
                step3_time = time.time() - t0
                logger.info(f"   ✅ Field decomposition complete in {step3_time:.1f}s")
                logger.info(f"   Complex fields decomposed: {len(helper_fields)}")
            else:
                logger.info("\n⏭️  STEP 3: Field decomposition skipped (disabled)")

            # ── STEP 4: Seed Prompt Generation (Claude) ───────────────
            logger.info("\n✍️  STEP 4: Seed Prompt Generation")
            t0 = time.time()
            seed_result = agent_generate_seed_prompt(
                document_analysis, decomposition_report, expected_values, api_key, tracker,
                domain_knowledge=domain_knowledge,
            )
            if not seed_result.success or not seed_result.parsed_json:
                raise RuntimeError(f"Seed prompt generation failed: {seed_result.error}")
            current_prompt_text = seed_result.parsed_json.get("prompt_text", "")
            current_prompt_set = seed_result.parsed_json.get("prompt_set", {})
            step4_time = time.time() - t0
            logger.info(f"   ✅ Seed prompt generated in {step4_time:.1f}s")

            audit.info("")
            audit.info("=" * 80)
            audit.info("SEED PROMPT (Step 4)")
            audit.info("=" * 80)
            audit.info(current_prompt_text)
            audit.info("-" * 40)
            audit.info(f"Prompt set: {json.dumps(current_prompt_set, indent=2)}")
            audit.info("=" * 80)

            # ── STEP 5: SIPDO Optimization Loop ──────────────────────
            logger.info("\n🔄 STEP 5: SIPDO Optimization Loop")
            logger.info(f"   Max iterations: {max_iterations}, Max difficulty: {difficulty_levels}")

            synthetic_docs = []
            all_test_data = []  # (document_text, expected_values) pairs for global check

            # Add original doc as first test data
            all_test_data.append((document_text, expected_values))

            for iteration in range(1, max_iterations + 1):
                difficulty = min(iteration, difficulty_levels)
                iter_start = time.time()

                logger.info(f"\n   ── Iteration {iteration}/{max_iterations} (difficulty {difficulty}) ──")

                # 5a. Generate synthetic data
                logger.info(f"   5a. Generating synthetic data (difficulty {difficulty})")
                synth_result = agent_generate_synthetic_data(
                    document_text, expected_values, difficulty, synthetic_docs, api_key, tracker
                )
                synth_doc = document_text  # fallback
                synth_expected = expected_values
                if synth_result.success and synth_result.parsed_json:
                    synth_doc = synth_result.parsed_json.get("synthetic_document", document_text)
                    synth_expected = synth_result.parsed_json.get("expected_values", expected_values)
                    synthetic_docs.append(synth_doc[:500])  # Store summary
                    all_test_data.append((synth_doc, synth_expected))

                audit.info("")
                audit.info(f"{'─' * 80}")
                audit.info(f"ITERATION {iteration}/{max_iterations} — Difficulty {difficulty}")
                audit.info(f"{'─' * 80}")
                audit.info(f"[SYNTHETIC DOCUMENT] ({len(synth_doc)} chars)")
                audit.info(synth_doc)
                audit.info(f"[SYNTHETIC EXPECTED VALUES]")
                audit.info(json.dumps(synth_expected, indent=2))

                # 5b. Evaluate prompt
                logger.info(f"   5b. Evaluating prompt (mode: {evaluation_mode})")
                if evaluation_mode == "coreplus":
                    extracted = _evaluate_via_coreplus(
                        current_prompt_set, file_content, filename
                    )
                else:
                    eval_result = evaluate_prompt_with_gpt(
                        current_prompt_text, current_prompt_set, synth_doc, tracker
                    )
                    extracted = eval_result.parsed_json or {} if eval_result.success else {}

                accuracy, failed_fields, field_details = _compute_accuracy(
                    synth_expected, extracted
                )
                logger.info(f"   Accuracy: {accuracy:.0%} — Failed: {failed_fields}")

                audit.info(f"[EVALUATION] Accuracy: {accuracy:.0%}")
                audit.info(f"[EVALUATION] Failed fields: {json.dumps(failed_fields)}")
                audit.info(f"[EVALUATION] Extracted: {json.dumps(extracted, indent=2, default=str)}")

                iter_action = "passed"

                # 5c. If errors, refine
                if failed_fields:
                    logger.info(f"   5c. Analyzing {len(failed_fields)} errors")
                    error_result = agent_analyze_errors(
                        current_prompt_text, synth_doc, synth_expected,
                        extracted, failed_fields, api_key, tracker,
                        domain_knowledge=domain_knowledge,
                    )
                    if error_result.success and error_result.parsed_json:
                        # 5d. Refine prompt
                        logger.info("   5d. Refining prompt")
                        refine_result = agent_refine_prompt(
                            current_prompt_text, current_prompt_set,
                            error_result.parsed_json, api_key, tracker
                        )
                        if refine_result.success and refine_result.parsed_json:
                            new_prompt_text = refine_result.parsed_json.get(
                                "refined_prompt_text", current_prompt_text
                            )
                            new_prompt_set = refine_result.parsed_json.get(
                                "refined_prompt_set", current_prompt_set
                            )

                            # 5e. Global verification — does revised prompt still work on all prior data?
                            logger.info("   5e. Global verification")
                            regression = False
                            for prev_doc, prev_expected in all_test_data[:-1]:
                                if evaluation_mode == "coreplus":
                                    prev_extracted = _evaluate_via_coreplus(
                                        new_prompt_set, file_content, filename
                                    )
                                else:
                                    prev_eval = evaluate_prompt_with_gpt(
                                        new_prompt_text, new_prompt_set, prev_doc, tracker
                                    )
                                    prev_extracted = (
                                        prev_eval.parsed_json or {} if prev_eval.success else {}
                                    )
                                prev_acc, prev_failed, _ = _compute_accuracy(
                                    prev_expected, prev_extracted
                                )
                                if prev_failed:
                                    logger.warning(
                                        f"   ⚠️  Regression on previous data: {prev_failed}"
                                    )
                                    regression = True
                                    break

                            if not regression:
                                current_prompt_text = new_prompt_text
                                current_prompt_set = new_prompt_set
                                iter_action = "refined"
                                logger.info("   ✅ Prompt refined successfully (no regression)")

                                audit.info(f"[REFINED PROMPT — Iteration {iteration}]")
                                audit.info(current_prompt_text)
                                audit.info(f"[REFINED PROMPT SET]")
                                audit.info(json.dumps(current_prompt_set, indent=2))
                            else:
                                iter_action = "regression_skipped"
                                logger.warning("   ⚠️  Refinement caused regression — keeping previous prompt")
                        else:
                            iter_action = "refine_failed"
                    else:
                        iter_action = "analysis_failed"
                else:
                    logger.info("   ✅ All fields correct at this difficulty")

                iter_time = time.time() - iter_start
                optimization_trace.append({
                    "iteration": iteration,
                    "difficulty": difficulty,
                    "accuracy": round(accuracy, 4),
                    "errors": failed_fields,
                    "action": iter_action,
                    "duration_seconds": round(iter_time, 2),
                })

                audit.info(f"[ITERATION {iteration} SUMMARY] action={iter_action} accuracy={accuracy:.0%} time={iter_time:.1f}s")

                # Early exit if perfect on max difficulty
                if accuracy == 1.0 and difficulty == difficulty_levels:
                    logger.info(f"   🎯 Perfect accuracy at max difficulty — early exit")
                    break

            logger.info(f"\n   SIPDO loop complete: {len(optimization_trace)} iterations")

            # ── STEP 6: Consistency Test ──────────────────────────────
            consistency_report = {}
            if consistency_test:
                logger.info("\n✅ STEP 6: Consistency Test")
                t0 = time.time()

                # Run prompt against original document
                if evaluation_mode == "coreplus":
                    final_extracted = _evaluate_via_coreplus(
                        current_prompt_set, file_content, filename
                    )
                else:
                    final_eval = evaluate_prompt_with_gpt(
                        current_prompt_text, current_prompt_set, document_text, tracker
                    )
                    final_extracted = final_eval.parsed_json or {} if final_eval.success else {}

                final_accuracy, final_failed, final_details = _compute_accuracy(
                    expected_values, final_extracted
                )

                # Run consistency auditor (Claude)
                audit_result = agent_audit_consistency(
                    current_prompt_text, document_text, expected_values,
                    final_extracted, api_key, tracker
                )
                if audit_result.success and audit_result.parsed_json:
                    consistency_report = audit_result.parsed_json
                else:
                    consistency_report = {
                        "overall_accuracy": final_accuracy,
                        "per_field": final_details,
                    }

                step6_time = time.time() - t0
                logger.info(f"   ✅ Consistency test complete in {step6_time:.1f}s")
                logger.info(f"   Final accuracy: {final_accuracy:.0%}")
            else:
                logger.info("\n⏭️  STEP 6: Consistency test skipped (disabled)")

            # ── STEP 7: Package Results ───────────────────────────────
            total_time = time.time() - start_time
            logger.info(f"\n📦 STEP 7: Packaging Results")
            logger.info(f"   Total time: {total_time:.1f}s")
            logger.info(f"   Token usage: {json.dumps(tracker.summary())}")

            audit.info("")
            audit.info("=" * 80)
            audit.info("FINAL OPTIMIZED PROMPT")
            audit.info("=" * 80)
            audit.info(current_prompt_text)
            audit.info("=" * 80)
            audit.info(f"FINAL PROMPT SET: {json.dumps(current_prompt_set, indent=2)}")
            if helper_fields:
                audit.info(f"HELPER FIELDS: {json.dumps(helper_fields, indent=2)}")
            if consistency_report:
                audit.info(f"CONSISTENCY REPORT: {json.dumps(consistency_report, indent=2, default=str)}")
            audit.info(f"TOTAL TIME: {total_time:.1f}s")
            audit.info(f"TOKEN USAGE: {json.dumps(tracker.summary(), indent=2)}")
            audit.info("=" * 80)
            audit.info("END OF AUDIT LOG")
            audit.info("=" * 80)

            # Clean up audit logger handlers
            for h in audit.handlers[:]:
                h.close()
                audit.removeHandler(h)

            return {
                "status": "success",
                "pipeline_id": pipeline_id,
                "document_name": filename,
                "optimized_prompt": current_prompt_text,
                "prompt_set": current_prompt_set,
                "helper_fields": helper_fields,
                "consistency_report": consistency_report,
                "optimization_trace": optimization_trace,
                "metadata": {
                    "total_iterations": len(optimization_trace),
                    "max_iterations": max_iterations,
                    "difficulty_levels": difficulty_levels,
                    "evaluation_mode": evaluation_mode,
                    "field_decomposition_applied": bool(helper_fields),
                    "consistency_test_run": consistency_test,
                    "total_time_seconds": round(total_time, 2),
                    "token_usage": tracker.summary(),
                    "estimated_cost_usd": round(tracker.estimated_cost, 4),
                },
                "processing_time_seconds": round(total_time, 2),
                "timestamp": datetime.utcnow().isoformat(),
            }

        except Exception as e:
            total_time = time.time() - start_time
            logger.error(f"Pipeline error: {e}", exc_info=True)
            audit.info(f"ERROR: {e}")
            audit.info(f"TOTAL TIME: {total_time:.1f}s")
            for h in audit.handlers[:]:
                h.close()
                audit.removeHandler(h)
            return {
                "status": "error",
                "pipeline_id": pipeline_id,
                "document_name": filename,
                "error": str(e),
                "optimization_trace": optimization_trace,
                "metadata": {
                    "total_time_seconds": round(total_time, 2),
                    "token_usage": tracker.summary(),
                },
                "processing_time_seconds": round(total_time, 2),
                "timestamp": datetime.utcnow().isoformat(),
            }


# ---------------------------------------------------------------------------
# Singleton accessor
# ---------------------------------------------------------------------------

_prompt_optimizer_service: Optional[PromptOptimizerService] = None


def get_prompt_optimizer_service() -> PromptOptimizerService:
    """Get or create the PromptOptimizer service singleton."""
    global _prompt_optimizer_service
    if _prompt_optimizer_service is None:
        _prompt_optimizer_service = PromptOptimizerService()
    return _prompt_optimizer_service
